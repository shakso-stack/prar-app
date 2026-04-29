"""
PRAR App — FastAPI Backend
Handles: Crossref/OpenAlex fetching, docx generation, Excel export
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import io
import re
import html
import requests
import time
import zipfile

app = FastAPI(title="PRAR Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

CROSSREF_BASE = "https://api.crossref.org/works"
CROSSREF_EMAIL = "prar.app@gmail.com"

# ─── Models ──────────────────────────────────────────────────────────────────

class JournalIssue(BaseModel):
    tier: str
    journal: str
    issn: str
    volume: str
    issue: Optional[str] = ""
    year: Optional[str] = ""

class FetchRequest(BaseModel):
    issues: List[JournalIssue]

class Article(BaseModel):
    author: str
    title: str
    journal: str
    volume: str
    issue: str
    year: str
    tier: str
    abstract: str
    link: str
    part: Optional[str] = ""

class GenerateRequest(BaseModel):
    installment_number: int
    season_year: str
    articles: List[Article]

# ─── Helpers ─────────────────────────────────────────────────────────────────

def _clean_num(value) -> str:
    if value is None or str(value).strip() == "":
        return ""
    try:
        return str(int(float(str(value).strip())))
    except (ValueError, TypeError):
        return str(value).strip()


def _clean_text(text: str) -> str:
    # Replace block-level tags with spaces to prevent word-merging
    text = re.sub(r'<(p|div|br|li|tr|td|th)(\s[^>]*)?>',  ' ', text, flags=re.IGNORECASE)
    text = re.sub(r'</(p|div|li|tr|td|th)>', ' ', text, flags=re.IGNORECASE)
    # Strip all remaining tags
    text = re.sub(r'<[^>]+>', '', text)
    # Decode HTML entities
    text = html.unescape(text)
    # Collapse multiple spaces/newlines into single space
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def parse_abstract(item: dict) -> str:
    raw = item.get("abstract", "")
    if not raw:
        return ""
    return _clean_text(raw)


def fetch_openalex_abstract(doi: str) -> str:
    if not doi:
        return ""
    try:
        url = f"https://api.openalex.org/works/https://doi.org/{doi}"
        resp = requests.get(url, params={"mailto": CROSSREF_EMAIL}, timeout=15)
        if resp.status_code != 200:
            return ""
        data = resp.json()
        inv_index = data.get("abstract_inverted_index")
        if not inv_index:
            return ""
        word_positions = []
        for word, positions in inv_index.items():
            for pos in positions:
                word_positions.append((pos, word))
        word_positions.sort(key=lambda x: x[0])
        return " ".join(w for _, w in word_positions)
    except Exception:
        return ""


def parse_authors(item: dict) -> str:
    authors = item.get("author", [])
    parts = []
    for a in authors:
        given = a.get("given", "")
        family = a.get("family", "")
        name = f"{given} {family}".strip() if given else family
        if name:
            parts.append(name)
    return ", ".join(parts)


def fetch_articles(issn: str, volume: str, issue: str) -> list:
    matched = []
    params = {
        "filter": f"issn:{issn}",
        "rows": 100,
        "mailto": CROSSREF_EMAIL,
    }
    offset = 0
    total = None

    while True:
        params["offset"] = offset
        try:
            resp = requests.get(CROSSREF_BASE, params=params, timeout=30)
            resp.raise_for_status()
        except Exception as e:
            print(f"  Error fetching {issn}: {e}")
            return matched

        data = resp.json().get("message", {})
        items = data.get("items", [])
        if total is None:
            total = data.get("total-results", 0)

        for item in items:
            item_vol = str(item.get("volume", "")).strip()
            item_issue = str(item.get("issue", "")).strip()
            vol_match = (item_vol == volume)
            issue_match = (not issue) or (item_issue == issue)
            if vol_match and issue_match:
                matched.append(item)

        offset += len(items)
        if not items or offset >= total:
            break
        time.sleep(0.5)

    return matched


# ─── Part mapping ─────────────────────────────────────────────────────────────

PART_MAP = {
    "Arab Law Quarterly": "Part 1",
    "Arab Media & Society": "Part 1",
    "Arab Studies Journal": "Part 1",
    "Arab Studies Quarterly": "Part 1",
    "British Journal of Middle Eastern Studies": "Part 1",
    "Comparative Studies of South Asia, Africa and the Middle East": "Part 1",
    "Contemporary Arab Affairs": "Part 1",
    "International Journal of Middle East Studies": "Part 1",
    "Israel Studies": "Part 1",
    "Journal of Middle East Women's Studies": "Part 1",
    "Journal of Palestine Studies": "Part 1",
    "Mediterranean Politics": "Part 1",
    "Middle East Critique": "Part 1",
    "Middle East Law and Governance": "Part 1",
    "Middle East Policy": "Part 1",
    "Middle East Quarterly": "Part 1",
    "Middle East Report": "Part 1",
    "Middle East Review of International Affairs": "Part 1",
    "The Middle East Journal": "Part 1",
    "Alif: Journal of Comparative Poetics": "Part 2",
    "Anatolian Studies": "Part 2",
    "Arabica": "Part 2",
    "Ars Orientalis": "Part 2",
    "British Museum Studies in Ancient Egypt and Sudan": "Part 2",
    "Bulletin of the American Society of Overseas Research": "Part 2",
    "Bulletin of the School of Oriental and African Studies": "Part 2",
    "Comparative Studies in Society and History": "Part 2",
    "Dead Sea Discoveries": "Part 2",
    "Iran": "Part 2",
    "Iran and the Caucasus": "Part 2",
    "Iranian Studies": "Part 2",
    "Iraq": "Part 2",
    "Islamic Law and Society": "Part 2",
    "Journal of Arabic Literature": "Part 2",
    "Journal of Arabic and Islamic Studies": "Part 2",
    "Journal of Contemporary History": "Part 2",
    "Journal of Cuneiform Studies": "Part 2",
    "Journal of Islamic Studies": "Part 2",
    "Journal of Near Eastern Studies": "Part 2",
    "Journal of Qur'anic Studies": "Part 2",
    "Journal of Social History": "Part 2",
    "Journal of the American Oriental Society": "Part 2",
    "Journal of the Economic and Social History of the Orient": "Part 2",
    "MELA Notes": "Part 2",
    "Mamluk Studies Review": "Part 2",
    "Middle Eastern Literatures": "Part 2",
    "Middle Eastern Studies": "Part 2",
    "Muqarnas": "Part 2",
    "Near Eastern Archaeology": "Part 2",
    "Oriens": "Part 2",
    "Proceedings of the Seminar for Arabian Studies": "Part 2",
    "Critical Studies on Terrorism": "Part 3",
    "Defense and Peace Economics": "Part 3",
    "Democratization": "Part 3",
    "Dynamics of Asymmetric Conflict": "Part 3",
    "Global Change Peace & Security": "Part 3",
    "Government and Opposition": "Part 3",
    "Journal of Peace Research": "Part 3",
    "Security Studies": "Part 3",
    "Studies in Conflict & Terrorism": "Part 3",
    "Studies in Conflict &Terrorism": "Part 3",
    "Terrorism and Political Violence": "Part 3",
    "Acta Politica": "Part 4",
    "American Anthropologist": "Part 4",
    "American Economic Review": "Part 4",
    "American Ethnologist": "Part 4",
    "American Journal of Political Science": "Part 4",
    "American Political Science Review": "Part 4",
    "Anthropology & Education Quarterly": "Part 4",
    "British Journal of Political Science": "Part 4",
    "Comparative Political Studies": "Part 4",
    "Comparative Politics": "Part 4",
    "Constellations": "Part 4",
    "Critical Review": "Part 4",
    "Development and Change": "Part 4",
    "Economic Affairs": "Part 4",
    "Economic Development Quarterly": "Part 4",
    "Economics and Politics": "Part 4",
    "Electoral Studies": "Part 4",
    "European Journal of Development Research": "Part 4",
    "European Journal of International Relations": "Part 4",
    "European Journal of Political Research": "Part 4",
    "European Political Science Review": "Part 4",
    "Georgetown Journal of International Affairs": "Part 4",
    "Global Media Journal": "Part 4",
    "International Affairs": "Part 4",
    "International Interactions": "Part 4",
    "International Journal of Urban and Regional Research": "Part 4",
    "International Organization": "Part 4",
    "International Political Science Review": "Part 4",
    "International Political Sociology": "Part 4",
    "International Politics": "Part 4",
    "International Relations": "Part 4",
    "International Studies": "Part 4",
    "International Studies Perspectives": "Part 4",
    "International Studies Quarterly": "Part 4",
    "International Studies Review": "Part 4",
    "Journal of Civil Society": "Part 4",
    "Journal of Democracy": "Part 4",
    "Journal of Developing Societies": "Part 4",
    "Journal of Economic Cooperation and Development": "Part 4",
    "Journal of Economic Literature": "Part 4",
    "Journal of Economic Policy Reform": "Part 4",
    "Journal of Globalization and Development": "Part 4",
    "Journal of Institutional Economics": "Part 4",
    "Journal of Political Economy": "Part 4",
    "Law & Development Review": "Part 4",
    "Millennium: Journal of International Studies": "Part 4",
    "New Global Studies": "Part 4",
    "New Left Review": "Part 4",
    "Oxford Development Studies": "Part 4",
    "PS: Political Science & Politics": "Part 4",
    "Participations: Journal of Audience & Reception Studies": "Part 4",
    "Perspectives on Politics": "Part 4",
    "Political Science Quarterly": "Part 4",
    "Political Studies": "Part 4",
    "Politics & Society": "Part 4",
    "Politics, Philosophy & Economics": "Part 4",
    "Polity": "Part 4",
    "Progress in Development Studies": "Part 4",
    "Review of African Political Economy": "Part 4",
    "Review of International Political Economy": "Part 4",
    "Review of Middle East Economics and Finance": "Part 4",
    "Review of Political Economy": "Part 4",
    "Review of Radical Political Economics": "Part 4",
    "Social & Legal Studies": "Part 4",
    "The American Journal of Economics and Sociology": "Part 4",
    "The Developing Economies": "Part 4",
    "The Economic Journal": "Part 4",
    "The European Journal of Development Research": "Part 4",
    "The Journal of Development Studies": "Part 4",
    "The Journal of Politics": "Part 4",
    "The Washington Quarterly": "Part 4",
    "Third World Quarterly": "Part 4",
    "World Politics": "Part 4",
}

PART_COLORS = {
    "Part 1": "#FF6600",
    "Part 2": "#2196F3",
    "Part 3": "#4CAF50",
    "Part 4": "#9C27B0",
}

# ─── Number to words ──────────────────────────────────────────────────────────

def number_to_words(n: int) -> str:
    ones = ["", "one", "two", "three", "four", "five", "six", "seven", "eight",
            "nine", "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen",
            "sixteen", "seventeen", "eighteen", "nineteen"]
    tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]
    if n == 0:
        return "zero"
    if n < 20:
        return ones[n]
    if n < 100:
        return tens[n // 10] + ("-" + ones[n % 10] if n % 10 else "")
    if n < 1000:
        rest = n % 100
        return ones[n // 100] + " hundred" + (" " + number_to_words(rest) if rest else "")
    return str(n)

def number_to_ordinal_words(n: int) -> str:
    irregular = {
        1: "first", 2: "second", 3: "third", 4: "fourth", 5: "fifth",
        6: "sixth", 7: "seventh", 8: "eighth", 9: "ninth", 10: "tenth",
        11: "eleventh", 12: "twelfth", 13: "thirteenth", 14: "fourteenth",
        15: "fifteenth", 16: "sixteenth", 17: "seventeenth", 18: "eighteenth",
        19: "nineteenth", 20: "twentieth", 21: "twenty-first", 22: "twenty-second",
        23: "twenty-third", 24: "twenty-fourth", 25: "twenty-fifth",
        26: "twenty-sixth", 27: "twenty-seventh", 28: "twenty-eighth",
        29: "twenty-ninth", 30: "thirtieth", 31: "thirty-first",
        32: "thirty-second", 33: "thirty-third", 34: "thirty-fourth",
        35: "thirty-fifth", 36: "thirty-sixth", 37: "thirty-seventh",
        38: "thirty-eighth", 39: "thirty-ninth", 40: "fortieth",
        50: "fiftieth", 60: "sixtieth", 70: "seventieth", 80: "eightieth",
        90: "ninetieth", 100: "one hundredth",
    }
    return irregular.get(n, number_to_words(n) + "th")

# ─── Routes ──────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/fetch")
def fetch(req: FetchRequest):
    all_articles = []
    results = []

    for entry in req.issues:
        journal = entry.journal.strip()
        issn = entry.issn.strip()
        volume = entry.volume.strip()
        issue = (entry.issue or "").strip()
        year = (entry.year or "").strip()
        tier = str(entry.tier).strip()

        items = fetch_articles(issn, volume, issue)
        filtered = [i for i in items if i.get("type") == "journal-article"]

        for item in filtered:
            doi = item.get("DOI", "")
            link = f"https://doi.org/{doi}" if doi else ""
            title_list = item.get("title", [])
            subtitle_list = item.get("subtitle", [])
            title = _clean_text(title_list[0]) if title_list else ""
            subtitle = _clean_text(subtitle_list[0]) if subtitle_list else ""
            if subtitle:
                title = f"{title}: {subtitle}"
            abstract = parse_abstract(item)
            if not abstract:
                abstract = fetch_openalex_abstract(doi)
            part = PART_MAP.get(journal, "")
            all_articles.append({
                "author": parse_authors(item),
                "title": title,
                "journal": journal,
                "volume": item.get("volume", volume),
                "issue": item.get("issue", issue),
                "year": year,
                "tier": tier,
                "abstract": abstract or "Not available",
                "link": link,
                "part": part,
            })

        results.append({"journal": journal, "count": len(filtered)})
        time.sleep(0.8)

    return {"articles": all_articles, "summary": results}


@app.post("/generate")
def generate(req: GenerateRequest):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    ordinal = number_to_ordinal_words(req.installment_number)
    intro_text = (
        f"This is the {ordinal} installment of the Peer-Reviewed Articles Review (PRAR), "
        f"a periodic review of peer-reviewed articles published in academic journals "
        f"relevant to the study of the Middle East and the broader Muslim world. "
        f"The articles listed below were selected based on their relevance to the region "
        f"and its peoples, and are arranged by journal title."
    )

    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        i_elem = OxmlElement("w:i")
        rPr.append(i_elem)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def set_font(run, size_pt=11, italic=False, color=None):
        run.font.name = "Garamond"
        run.font.size = Pt(size_pt)
        run.font.italic = italic
        run.font.bold = False
        if color:
            run.font.color.rgb = RGBColor(*color)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for part_num in range(1, 5):
            part_name = f"Part {part_num}"
            part_articles = [a for a in req.articles if a.part == part_name]
            part_articles.sort(key=lambda a: a.journal.lower())

            doc = Document()

            # Page margins
            from docx.shared import Inches
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

            # Remove default styles bold
            style = doc.styles["Normal"]
            style.font.name = "Garamond"
            style.font.size = Pt(11)
            style.font.bold = False

            # Document title
            doc_title = f"Peer-Reviewed Articles Review: {req.season_year} ({part_name})"
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_after = Pt(6)
            title_run = title_para.add_run(doc_title)
            title_run.font.name = "Garamond"
            title_run.font.size = Pt(14)
            title_run.font.bold = False
            title_run.font.italic = True

            # Intro paragraph
            intro_para = doc.add_paragraph()
            intro_para.paragraph_format.space_after = Pt(12)
            intro_run = intro_para.add_run(intro_text)
            set_font(intro_run)

            # Articles
            seen_journals = set()
            for article in part_articles:
                # Journal header (italic, only first occurrence)
                journal_display = f"{article.journal} (Volume {article.volume}, Issue {article.issue})"
                if article.journal not in seen_journals:
                    j_para = doc.add_paragraph()
                    j_para.paragraph_format.space_before = Pt(10)
                    j_para.paragraph_format.space_after = Pt(2)
                    j_run = j_para.add_run(journal_display)
                    set_font(j_run, size_pt=11, italic=True)
                    seen_journals.add(article.journal)

                # Title as hyperlink
                title_para = doc.add_paragraph()
                title_para.paragraph_format.space_before = Pt(8)
                title_para.paragraph_format.space_after = Pt(2)
                if article.link:
                    add_hyperlink(title_para, article.title, article.link)
                else:
                    t_run = title_para.add_run(article.title)
                    set_font(t_run)

                # Author
                author_para = doc.add_paragraph()
                author_para.paragraph_format.space_after = Pt(2)
                author_run = author_para.add_run(f"By: {article.author}")
                set_font(author_run)

                # Abstract
                abstract_para = doc.add_paragraph()
                abstract_para.paragraph_format.space_after = Pt(10)
                abstract_run = abstract_para.add_run(f"Abstract: {article.abstract}")
                set_font(abstract_run)

            fname = f"PRAR_Installment_{req.installment_number}_{part_name.replace(' ', '_')}.docx"
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            zf.writestr(fname, doc_buffer.getvalue())

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=PRAR_Installment_{req.installment_number}.zip"}
    )


@app.post("/export-excel")
def export_excel(req: GenerateRequest):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    COLUMNS = ["Author", "Title", "Journal", "Volume", "Issue", "Year", "Tier", "Abstract", "Link", "Part"]
    HEADER_FILL = PatternFill("solid", start_color="1A237E", end_color="1A237E")
    HEADER_FONT = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    BODY_FONT = Font(name="Arial", size=10)
    THIN = Border(
        left=Side(style="thin", color="DDDDDD"), right=Side(style="thin", color="DDDDDD"),
        top=Side(style="thin", color="DDDDDD"), bottom=Side(style="thin", color="DDDDDD"),
    )
    PART_FILL_COLORS = {
        "Part 1": "FFE0CC", "Part 2": "BBDEFB",
        "Part 3": "C8E6C9", "Part 4": "E1BEE7",
    }

    sheets_data = {"All Articles": req.articles}
    for p in ["Part 1", "Part 2", "Part 3", "Part 4"]:
        sheets_data[p] = [a for a in req.articles if a.part == p]

    for sheet_name, articles in sheets_data.items():
        ws = wb.create_sheet(sheet_name)
        for ci, col in enumerate(COLUMNS, 1):
            cell = ws.cell(row=1, column=ci, value=col)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = THIN
        ws.row_dimensions[1].height = 24
        ws.freeze_panes = "A2"

        sorted_articles = sorted(articles, key=lambda a: a.journal.lower())
        for ri, art in enumerate(sorted_articles, 2):
            row_data = [art.author, art.title, art.journal, art.volume,
                        art.issue, art.year, art.tier, art.abstract, art.link, art.part]
            part_color = PART_FILL_COLORS.get(art.part, "FFFFFF")
            for ci, val in enumerate(row_data, 1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = BODY_FONT
                cell.border = THIN
                cell.alignment = Alignment(vertical="top", wrap_text=(ci in (2, 3, 8)))
                if ci == 3:  # journal column colored by part
                    cell.fill = PatternFill("solid", start_color=part_color, end_color=part_color)
            ws.row_dimensions[ri].height = 50 if row_data[7] else 18

        COL_WIDTHS = [30, 50, 30, 8, 8, 8, 8, 80, 40, 10]
        for ci, w in enumerate(COL_WIDTHS, 1):
            ws.column_dimensions[get_column_letter(ci)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=PRAR_Installment_{req.installment_number}_Compiled.xlsx"}
    )
